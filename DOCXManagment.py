from docx import Document

#uses list to fill in blanks for report
def create_report(device):
    document = Document()
    document.add_paragraph('On ' + device[0] + ', I, Detective ' + device[2] + ' was given the following device by '
                           + 'Detective ' + device[1] + ', ' + 'in ' + 'reference to '
                           + device[6] + ' case number ' + device[9] + '.')
    document.add_paragraph(device[7] + " " + " Serial: " + device[8])
    document.add_paragraph("The device was connected to a forensic workstation and processed using Cellebrite's UFED "
                           "4PC and Physical Analyzer. " "A report was generated and was provided to Detective " +
                           device[1] + ". This concludes my involvement in this investigation.")
    document.save(device[9] + '.docx')
